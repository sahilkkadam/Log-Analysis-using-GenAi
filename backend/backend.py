import os
import shutil
import json
import hashlib
from fastapi import FastAPI, File, UploadFile
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import chromadb
import ollama
from typing import List, Generator
from pydantic import BaseModel
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.document_loaders.csv_loader import CSVLoader
import pandas as pd
import docx

# Class model for the request body
class Data(BaseModel):
    """
    Model to define the request body for asking questions.

    Attributes:
        question (str): The question to be answered.
        file_names (List[str]): List of file names to search for answers.
    """
    question: str 
    file_names: List[str] 


# Initialize FastAPI app
app = FastAPI()

# Allow CORS for all origins (development purposes)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Function to create necessary directories
def create_directory(directory_path):
    """
    Creates a directory if it doesn't exist.

    Args:
        directory_path (str): Path of the directory to be created.
    """
    try:
        os.makedirs(directory_path, exist_ok=True)
        print(f"Directory created at: {directory_path}")
    except OSError as e:
        print(f"Error creating directory: {e}")

# Save uploaded files locally
@app.post("/upload-file/")
async def upload_file(file: UploadFile = File(...)):
    """
    Uploads a file and saves it locally.

    Args:
        file (UploadFile): The file to be uploaded.

    Returns:
        dict: Contains the filename and file path.
    """
    try:
        upload_directory = 'uploads'
        create_directory(upload_directory)
        file_path = os.path.join(upload_directory, file.filename)
        
        # Save the uploaded file to the specified directory
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        return {"filename": file.filename, "file_path": file_path}
    except Exception as e:
        return {"error": str(e)}

# Function to extract text from a PDF file
def extract_text_from_pdf(pdf_path):
    """
    Extracts text from a PDF file.

    Args:
        pdf_path (str): Path to the PDF file.

    Returns:
        str: Extracted text from the PDF file.
    """
    try:
        loader = PyPDFLoader(pdf_path)
        pages = loader.load_and_split()
        text = "\n".join(page.page_content for page in pages)
        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return ""

# Function to extract text from a DOCX file
def extract_text_from_docx(docx_path):
    """
    Extracts text from a DOCX file.

    Args:
        docx_path (str): Path to the DOCX file.

    Returns:
        str: Extracted text from the DOCX file.
    """
    try:
        doc = docx.Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""

# Function to extract text from a TXT file
def extract_text_from_txt(txt_path):
    """
    Extracts text from a TXT file.

    Args:
        txt_path (str): Path to the TXT file.

    Returns:
        str: Extracted text from the TXT file.
    """
    try:
        loader = TextLoader(txt_path)
        documents = loader.load()
        text = "\n".join(doc.page_content for doc in documents)
        return text
    except Exception as e:
        print(f"Error extracting text from TXT: {e}")
        return ""

# Function to extract text from a CSV file
def extract_text_from_csv(csv_path):
    """
    Extracts text from a CSV file.

    Args:
        csv_path (str): Path to the CSV file.

    Returns:
        str: Extracted text from the CSV file.
    """
    try:
        loader = CSVLoader(file_path=csv_path)
        data = loader.load()
        text = "\n".join([str(record) for record in data])
        return text
    except Exception as e:
        print(f"Error extracting text from CSV: {e}")
        return ""

# Function to extract text from an XLSX file
def extract_text_from_xlsx(xlsx_path):
    """
    Extracts text from an XLSX file.

    Args:
        xlsx_path (str): Path to the XLSX file.

    Returns:
        str: Extracted text from the XLSX file.
    """
    try:
        df = pd.read_excel(xlsx_path)
        text = df.to_string(index=False)
        return text
    except Exception as e:
        print(f"Error extracting text from XLSX: {e}")
        return ""

# Function to get file extension
def get_file_extension(file_path):
    """
    Gets the extension of a file.

    Args:
        file_path (str): Path to the file.

    Returns:
        str: File extension in lowercase.
    """
    try:
        return os.path.splitext(file_path)[1].lower()
    except Exception as e:
        print(f"Error getting file extension: {e}")
        return ""

# Function to extract text based on file type
def extract_text(file_path):
    """
    Extracts text from a file based on its extension.

    Args:
        file_path (str): Path to the file.

    Returns:
        str: Extracted text from the file.
    """
    try:
        ext = get_file_extension(file_path)
        if ext == '.pdf':
            return extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return extract_text_from_docx(file_path)
        elif ext == '.txt':
            return extract_text_from_txt(file_path)
        elif ext == '.csv':
            return extract_text_from_csv(file_path)
        elif ext == '.xlsx':
            return extract_text_from_xlsx(file_path)
        else:
            raise ValueError("Unsupported file type")
    except Exception as e:
        print(f"Error extracting text: {e}")
        return ""

# Function to chunk text into smaller pieces
def chunk_text(text, chunk_size=1000, chunk_overlap=50):
    """
    Splits text into smaller chunks for better processing.

    Args:
        text (str): The text to be split.
        chunk_size (int): The size of each chunk.
        chunk_overlap (int): The overlap between chunks.

    Returns:
        List[str]: List of text chunks.
    """
    try:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        doc_output = splitter.split_documents([Document(page_content=text)])
        # print(doc_output)
        # Convert the Document objects to a list of strings
        result = [doc.page_content for doc in doc_output]
        # print(result)
        return result
    except Exception as e:
        print(f"Error chunking text: {e}")
        return []

# Store embeddings in ChromaDB
def chromadb_vector_store(embeddings, paragraphs, collection_name):
    """
    Stores embeddings in a ChromaDB collection.

    Args:
        embeddings (List): List of embeddings to store.
        paragraphs (List[str]): List of text paragraphs corresponding to the embeddings.
        collection_name (str): Name of the ChromaDB collection.

    Returns:
        chromadb.Collection: The collection where embeddings are stored.
    """
    try:
        client = chromadb.HttpClient(host='localhost', port=8001)  # ChromaDB port
        collection = client.get_or_create_collection(
            name=collection_name, metadata={"hnsw:space": "cosine"})

        # Add embeddings to the collection
        n = len(paragraphs)
        collection.add(
            ids=[str(id) for id in range(n)],
            embeddings=[embedding for embedding in embeddings],
            documents=[paragraph for paragraph in paragraphs],
            metadatas=[{"doc_id": i} for i in range(n)],
        )

        print("Stored embeddings in ChromaDB collection")
        return collection
    except Exception as e:
        print(f"Error storing embeddings in ChromaDB: {e}")
        return None
    
# Generate a hash from an input string
def generate_hash(input_string):
    """
    Generates a SHA-256 hash from a string.

    Args:
        input_string (str): The input string to hash.

    Returns:
        str: The resulting hash.
    """
    try:
        hash_object = hashlib.sha256(input_string.encode())
        return hash_object.hexdigest()
    except Exception as e:
        print(f"Error generating hash: {e}")
        return ""

# Hash map to keep track of file metadata
file_hash_map = {}

# Add file metadata to the hash map
def add_to_hash_map(file_name):
    """
    Adds file metadata to the hash map by generating a hash.

    Args:
        file_name (str): The name of the file.
    """
    try:
        file_hash_map[file_name] = generate_hash(file_name)
    except Exception as e:
        print(f"Error adding to hash map: {e}")

# Load and save embeddings using JSON files
def save_embeddings(filename, embeddings):
    """
    Save embeddings to a JSON file.

    Parameters:
    filename (str): The name of the file to save the embeddings to.
    embeddings (List[List[float]]): The embeddings to save.
    """
    try:
        if not os.path.exists("embeddings"):
            os.makedirs("embeddings")
        with open(f"embeddings/{filename}.json", "w") as f:
            json.dump(embeddings, f)
    except Exception as e:
        print(f"Error saving embeddings: {e}")

def load_embeddings(filename):
    """
    Load embeddings from a JSON file.

    Parameters:
    filename (str): The name of the file to load the embeddings from.

    Returns:
    List[List[float]]: The loaded embeddings, or False if loading fails.
    """
    try:
        if not os.path.exists(f"embeddings/{filename}.json"):
            return False
        with open(f"embeddings/{filename}.json", "r") as f:
            return json.load(f)
    except Exception as e:
        print(f"Error loading embeddings: {e}")
        return False

# Function to get or generate embeddings
def get_embeddings(filename, modelname, chunks):
    """
    Get or generate embeddings for the provided chunks of text.

    Parameters:
    filename (str): The name of the file associated with the embeddings.
    modelname (str): The name of the model to use for generating embeddings.
    chunks (List[str]): The chunks of text to generate embeddings for.

    Returns:
    List[List[float]]: The embeddings for the provided chunks.
    """
    try:
        if (embeddings := load_embeddings(filename)) is not False:
            return embeddings
        embeddings = [
            ollama.embeddings(model=modelname, prompt=chunk)["embedding"]
            for chunk in chunks
        ]
        save_embeddings(filename, embeddings)
        return embeddings
    except Exception as e:
        print(f"Error getting embeddings: {e}")
        return []

# Function to combine results and pick top 7 chunks
def combine_and_select_top_chunks(results_list, top_n=7):
    """
    Combine results from multiple collections and select the top N chunks based on similarity.

    Parameters:
    results_list (List[dict]): A list of results from different collections.
    top_n (int): The number of top chunks to select. Default is 7.

    Returns:
    List[str]: The top N chunks of text.
    """
    try:    
        combined_results = []
        for result in results_list:
            distances = result.get("distances", [])[0]
            documents = result.get("documents", [])[0]
            combined_results.extend(zip(distances, documents))
        
        # Sort combined results by distance (similarity score)
        combined_results.sort(key=lambda x: x[0])
        
        # Select top N results
        top_chunks = [doc for _, doc in combined_results[:top_n]]
        return top_chunks
    except Exception as e:
        print(f"Error combining and selecting top chunks: {e}")
        return []
